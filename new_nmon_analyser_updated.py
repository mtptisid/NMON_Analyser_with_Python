#!/usr/bin/env python3
###########################################################################
######## Script for nmon anyalysis for linux servers                     ##
######## Created By    : MATHAPATI Siddharamayya                         ##
######## Creation Date : 1st December 2024                              ##
######## Email         : msidrm455@gmail.com                             ##
######## Version       : 2.0                                             ##
###########################################################################

import os
import re
import shutil
import smtplib
import zipfile
import paramiko
import argparse
import logging
from docx import Document
from io import BytesIO
import numpy as np
import pandas as pd
from lxml import etree
from scp import SCPClient
from docx import Document
from email import encoders
from shutil import copyfile
from docx.oxml.ns import qn
from itertools import product
from docx.shared import Inches, Pt
from docx.oxml import parse_xml
import matplotlib.pyplot as plt
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from email.mime.text import MIMEText
from datetime import datetime, timedelta
from email.mime.base import MIMEBase
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from email.mime.multipart import MIMEMultipart


def ssh_copy_file(host, remote_path, local_path):
    """Connects via SSH to copy a file, matching any extension."""
    try:
        # Extract the base path and file name without extension
        base_path, file_name = os.path.split(remote_path)
        file_base, _ = os.path.splitext(file_name)

        # Connect to the remote host
        ssh = paramiko.SSHClient()
        ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
        ssh.connect(hostname=host, username='sysadm', port=77)

        # List files in the remote directory
        stdin, stdout, stderr = ssh.exec_command(f'ls {base_path}')
        files = stdout.read().decode().splitlines()

        # Find matching files
        matching_files = [f for f in files if f.startswith(f"{base_path}/{file_base}")]

        if not matching_files:
            print(f"No matching files found for {remote_path}")
            ssh.close()
            return

        # Copy each matching file
        with SCPClient(ssh.get_transport()) as scp:
            for remote_file in matching_files:
                # Construct the local file path
                local_file = os.path.join(local_path, os.path.basename(remote_file))
                print(f"Copying {remote_file} to {local_file}")
                scp.get(remote_file, local_file)

        ssh.close()
    except Exception as e:
        print(f"Error during SSH/SCP: {e}")
        exit(1)


def process_system_data(file_path, cpu_prefix="CPU", sections=None, timestamp_column="Timestamp"):
    """
    Processes system data from a file and organizes it into DataFrames for CPU and other sections.

    :param file_path: Path to the file containing system data.
    :param cpu_prefix: Prefix used for CPU data (e.g., "CPU").
    :param sections: List of section prefixes to process (e.g., ["MEM", "VM", "PROC"]).
    :param timestamp_column: The name for the timestamp column to include.
    :return: Dictionary of DataFrames keyed by section prefixes, including CPU data.
    """
    global CONFIG

    if sections is None:
        sections = ["CPU_ALL", "MEM", "VM", "PROC", "NET", "TOP", "NETPACKET", "JFSFILE",
                    "DISKBUSY", "DISKREAD", "DISKWRITE", "DISKXFER", "DISKBSIZE"]

    # Dictionaries to hold the data
    data_dict = {section: {"columns": [], "data": []} for section in [cpu_prefix] + sections}
    found_headers = set()

    with open(file_path, 'r') as file:
        top_section_found = False  # Flag to indicate when we encounter the top section
        for line in file:
            line = line.strip()

            # Skip empty lines or lines that don't start with a valid section
            if not line or not any(line.startswith(section) for section in [cpu_prefix] + sections):
                continue

            # Check for TOP section
            if line.startswith("TOP"):
                # Skip the first occurrence (header line with %CPU Utilisation)
                if not top_section_found:
                    top_section_found = True
                    continue  # Skip the first header line

                # The second line is the actual column header for TOP
                if "TOP" not in found_headers:
                    parts = re.split(r",\s*", line)
                    data_dict["TOP"]["columns"] = [timestamp_column] + parts[1:]  # Exclude 'TOP' as it's redundant
                    found_headers.add("TOP")
                else:
                    # Process data lines under TOP
                    parts = re.split(r",\s*", line)
                    if len(parts) == len(data_dict["TOP"]["columns"]):
                        data_dict["TOP"]["data"].append(parts)
                    continue

            # Check for CPU data
            if line.startswith(cpu_prefix):
                parts = re.split(r",\s*", line)

                cpu_id = parts[0]  # The first part (e.g., CPU001) is used as the key
                if cpu_id not in data_dict:
                    data_dict[cpu_id] = {"columns": [], "data": []}

                # If header not found for this CPU, extract headers
                if cpu_id not in found_headers:
                    data_dict[cpu_id]["columns"] = [cpu_prefix, timestamp_column] + parts[2:]
                    found_headers.add(cpu_id)
                else:
                    # Add data line, ensure it matches the header length
                    if len(parts) == len(data_dict[cpu_id]["columns"]):
                        data_dict[cpu_id]["data"].append(parts)
                continue

            # Check for other sections
            for section in sections:
                if line.startswith(section):
                    parts = re.split(r",\s*", line)

                    # If header not found for this section, extract headers
                    if section not in found_headers:
                        # Use the current line as the header for this section
                        data_dict[section]["columns"] = [section, timestamp_column] + parts[2:]
                        found_headers.add(section)
                    else:
                        # Add data line, ensure it matches the header length
                        if len(parts) == len(data_dict[section]["columns"]):
                            data_dict[section]["data"].append(parts)
                    break

    # Convert each section's data into a DataFrame
    dataframes = {}
    for section, data in data_dict.items():
        if data["data"]:  # Only create DataFrame if there is data
            # Ensure no mismatch in number of columns before creating DataFrame
            max_columns = len(data["columns"])
            valid_data = [row for row in data["data"] if len(row) == max_columns]

            # If there is valid data, create DataFrame
            if valid_data:
                df = pd.DataFrame(valid_data, columns=data["columns"])
                dataframes[section] = df

    return dataframes


def plot_cpu(cpu_all_df, start_time="00:00:00", interval_seconds=30, num_records=2879):
    """
    Plots Idle% and Busy% against timestamps for the CPU_ALL DataFrame.

    :param cpu_all_df: DataFrame containing the aggregated CPU_ALL data.
    :param start_time: Start time of the data collection (HH:MM:SS format).
    :param interval_seconds: Interval in seconds between data points (30 seconds for nmon).
    :param num_records: Total number of records in the dataset.
    """

    global CONFIG
    # Convert the relevant columns to numeric, forcing errors to NaN (in case of any non-numeric values)
    cpu_all_df['User%'] = pd.to_numeric(cpu_all_df['User%'], errors='coerce')
    cpu_all_df['Sys%'] = pd.to_numeric(cpu_all_df['Sys%'], errors='coerce')
    cpu_all_df['Wait%'] = pd.to_numeric(cpu_all_df['Wait%'], errors='coerce')
    cpu_all_df['Idle%'] = pd.to_numeric(cpu_all_df['Idle%'], errors='coerce')
    cpu_all_df['Steal%'] = pd.to_numeric(cpu_all_df['Steal%'], errors='coerce')

    # Generate a datetime range based on the number of records and interval
    start_datetime = datetime.strptime(start_time, "%H:%M:%S")
    cpu_all_df['Timestamp'] = [
        start_datetime + timedelta(seconds=i * interval_seconds) for i in range(num_records)
    ]

    # Calculate Busy% as the sum of User%, Sys%, and Wait% (but make sure it's <= 100%)
    cpu_all_df['Busy%'] = cpu_all_df['User%'] + cpu_all_df['Sys%']
    cpu_all_df['Busy%'] = cpu_all_df['Busy%'].clip(upper=100)  # Ensure it doesn't go over 100%

    # Set up the plot
    plt.figure(figsize=(15, 8))

    timestamp_col = "Timestamp"  # Replace with your timestamp column name
    columns_to_plot = ['User%', 'Sys%', 'Wait%', 'Idle%', 'Busy%']  # Replace with your column names

    # Colors for the stacked areas
    area_colors = ['blue', 'red', 'gray', 'orange', 'purple']

    # Plotting
    plt.figure(figsize=(18, 6))  # Set the figure size
    plt.stackplot(cpu_all_df[timestamp_col],
              [cpu_all_df[col] for col in columns_to_plot],
              labels=columns_to_plot,
              colors=area_colors,
              alpha=0.8)  # Adjust transparency for better visual clarity

    # Set x-axis ticks to 30-minute intervals
    x_ticks = pd.date_range(start=start_datetime, periods=num_records // 60, freq="30min")
    plt.xticks(x_ticks, [t.strftime("%H:%M:%S") for t in x_ticks], rotation=45)

    # Add labels, title, legend, and grid
    plt.xlabel('Time', fontsize=12)
    plt.ylabel('Percentage (%)', fontsize=12)
    plt.title(f'CPU Total {CONFIG["HOSTNAME"]} - {CONFIG["DATE"]}', fontsize=16)
    plt.legend(loc='upper right')
    #plt.grid(True, which='both', linestyle='--', linewidth=0.5)

    # Limit y-axis to 0-100% and set ticks at intervals of 20
    plt.ylim(0, 100)
    plt.yticks(range(0, 101, 10))  # Set y-axis ticks at 0, 20, 40, ..., 100%

    # Save plot as PNG
    plt.tight_layout()
    png_file = os.path.join(CONFIG["HOSTOUTDIR"], f'{CONFIG["HOSTNAME"]}_CPU_ALL.png')
    plt.savefig(png_file)
    plt.close()


def plot_memory(mem_dfs, start_time="00:00:00", interval_seconds=30, num_records=2879):
    """
    Plot a stacked area chart from a DataFrame.

    Parameters:
    - df (DataFrame): Input DataFrame containing the data.
    - timestamp_col (str): Column name for the x-axis (timestamps).
    - metrics (list): List of column names to plot as stacked areas.
    - colors (list): List of colors for each metric.
    - title (str): Title of the chart.
    - y_label (str): Label for the y-axis.
    """

    global CONFIG

    mem_dfs['memtotal'] = pd.to_numeric(mem_dfs['memtotal'], errors='coerce')
    mem_dfs['memfree'] = pd.to_numeric(mem_dfs['memfree'], errors='coerce')


    start_datetime = datetime.strptime(start_time, "%H:%M:%S")
    mem_dfs['Timestamp'] = [
        start_datetime + timedelta(seconds=i * interval_seconds) for i in range(num_records)
    ]

    mem_dfs['Memory Used'] = mem_dfs['memtotal'] - mem_dfs['memfree']

    timestamp_col = "Timestamp"  # Replace with your timestamp column name
    columns_to_plot = ['Memory Used', 'memfree']  # Replace with your column names

    area_colors = ['orange', 'cyan']

    # Ensure the timestamp column and metric columns exist in the DataFrame
    if timestamp_col not in mem_dfs.columns or not all(col in mem_dfs.columns for col in columns_to_plot):
        raise ValueError("Invalid column names provided for plotting.")

    # Plot the stacked area chart
    fig, ax = plt.subplots(figsize=(18, 6))
    plt.stackplot(mem_dfs[timestamp_col],
                  [mem_dfs[col] for col in columns_to_plot],
                  labels=columns_to_plot,
                  colors=area_colors,
                  alpha=0.8)

    # Set x-axis ticks to 30-minute intervals
    x_ticks = pd.date_range(start=start_datetime, periods=num_records // 60, freq="30min")
    plt.xticks(x_ticks, [t.strftime("%H:%M:%S") for t in x_ticks], rotation=45)

    # Add labels and title
    ax.set_xlabel('Time', fontsize=12)
    ax.set_ylabel('Memory size', fontsize=12)
    ax.set_title(f'Memory Total {CONFIG["HOSTNAME"]} - {CONFIG["DATE"]}', fontsize=14, pad=20)
    ax.tick_params(axis='x', rotation=45)

    # Adjust the layout and add padding
    plt.subplots_adjust(left=0.1, right=0.85, top=0.9, bottom=0.2)  # Adjust plot padding

    # Position the legend outside the plot
    ax.legend(
        loc='lower center',
        bbox_to_anchor=(1, 0.5),  # Place legend outside on the right margin
        title='Memory Metrics',
        fontsize=8,
        title_fontsize=10,
    )

    # Save plot as PNG
    plt.tight_layout()
    png_file = os.path.join(CONFIG["HOSTOUTDIR"], f'{CONFIG["HOSTNAME"]}_MEMORY.png')
    plt.savefig(png_file)
    plt.close()


def plot_swap(mem_dfs, start_time="00:00:00", interval_seconds=30, num_records=2879):
    """
    Plot a stacked area chart from a DataFrame.

    Parameters:
    - df (DataFrame): Input DataFrame containing the data.
    - timestamp_col (str): Column name for the x-axis (timestamps).
    - metrics (list): List of column names to plot as stacked areas.
    - colors (list): List of colors for each metric.
    - title (str): Title of the chart.
    - y_label (str): Label for the y-axis.
    """

    global CONFIG

    mem_dfs['swaptotal'] = pd.to_numeric(mem_dfs['swaptotal'], errors='coerce')
    mem_dfs['swapfree'] = pd.to_numeric(mem_dfs['swapfree'], errors='coerce')


    start_datetime = datetime.strptime(start_time, "%H:%M:%S")
    mem_dfs['Timestamp'] = [
        start_datetime + timedelta(seconds=i * interval_seconds) for i in range(num_records)
    ]

    mem_dfs['Swap Used'] = mem_dfs['swaptotal'] - mem_dfs['swapfree']

    timestamp_col = "Timestamp"  # Replace with your timestamp column name
    columns_to_plot = ['Swap Used', 'swapfree']  # Replace with your column names

    area_colors = ['red', 'green']

    # Ensure the timestamp column and metric columns exist in the DataFrame
    if timestamp_col not in mem_dfs.columns or not all(col in mem_dfs.columns for col in columns_to_plot):
        raise ValueError("Invalid column names provided for plotting.")

    # Plot the stacked area chart
    fig, ax = plt.subplots(figsize=(18, 6))
    plt.stackplot(mem_dfs[timestamp_col],
                  [mem_dfs[col] for col in columns_to_plot],
                  labels=columns_to_plot,
                  colors=area_colors,
                  alpha=0.8)

    # Set x-axis ticks to 30-minute intervals
    x_ticks = pd.date_range(start=start_datetime, periods=num_records // 60, freq="30min")
    plt.xticks(x_ticks, [t.strftime("%H:%M:%S") for t in x_ticks], rotation=45)

    # Add labels and title
    ax.set_xlabel('Time', fontsize=12)
    ax.set_ylabel('Swap size', fontsize=12)
    ax.set_title(f'Swap Total  {CONFIG["HOSTNAME"]} - {CONFIG["DATE"]}', fontsize=14, pad=20)
    ax.tick_params(axis='x', rotation=45)

    # Adjust the layout and add padding
    plt.subplots_adjust(left=0.1, right=0.85, top=0.9, bottom=0.2)  # Adjust plot padding

    # Position the legend outside the plot
    ax.legend(
        loc='lower center',
        bbox_to_anchor=(1, 0.5),  # Place legend outside on the right margin
        title='Swap Metrics',
        fontsize=8,
        title_fontsize=10,
    )

    # Save plot as PNG
    plt.tight_layout()
    png_file = os.path.join(CONFIG["HOSTOUTDIR"], f'{CONFIG["HOSTNAME"]}_SWAP.png')
    plt.savefig(png_file)
    plt.close()


def plot_filesystem(fs_df):
    """
    Plot the average values of JSPFILE data columns (excluding timestamp)
    as percentages with X-axis as column headers and Y-axis from 0 to 100%.

    Parameters:
    - df (DataFrame): Input DataFrame containing the data.
    - timestamp_col (str): Column name for the timestamp to exclude.
    - title (str): Title of the chart.
    - y_label (str): Label for the y-axis.
    """
    global CONFIG
    timestamp_col = ["Timestamp","JFSFILE"]
    for col in fs_df.columns:
        if col != timestamp_col:
            fs_df[col] = pd.to_numeric(fs_df[col], errors='coerce')

    columns_to_plot = [
    col for col in fs_df.columns if col != timestamp_col and pd.api.types.is_numeric_dtype(fs_df[col])
    ]

    if not columns_to_plot:
        raise ValueError("No numeric columns found to plot!")

    averages = fs_df[columns_to_plot].mean().dropna()

    colors = plt.cm.tab20(np.linspace(0, 1, len(averages)))
    # Exclude the timestamp column and calculate averages for other columns

    # Create the plot
    fig, ax = plt.subplots(figsize=(18, 8))
    bars = ax.bar(
        averages.index,
        averages.values,
        color=colors,
        edgecolor='black',
        alpha=0.8,
    )

    # Annotate each bar with its value
    for bar in bars:
        height = bar.get_height()
        ax.text(
            bar.get_x() + bar.get_width() / 2,
            height + 1,
            f"{height:.2f}%",
            ha='center',
            va='bottom',
            fontsize=10,
        )

    # Add labels, title, and ticks
    ax.set_xlabel('Filesytems', fontsize=12)
    ax.set_ylabel('Usages(%) ', fontsize=12)
    ax.set_title(f'Filesytems Usage  {CONFIG["HOSTNAME"]} - {CONFIG["DATE"]}', fontsize=14, pad=20)
    ax.set_ylim(0, 100)  # Set Y-axis from 0 to 100%
    ax.set_yticks(range(0, 101, 10))  # Tick every 10%
    ax.tick_params(axis='x', rotation=45)  # Rotate X-axis labels for readability

    ax.set_xticks(range(len(averages.index)))
    ax.set_xticklabels(averages.index, rotation=45, ha='right', fontsize=10)

    # Save plot as PNG
    plt.tight_layout()
    png_file = os.path.join(CONFIG["HOSTOUTDIR"], f'{CONFIG["HOSTNAME"]}_FileSystem.png')
    plt.savefig(png_file)
    plt.close()


def plot_network_usage(net_df, start_time="00:00:00", interval_seconds=30, num_records=2879):
    """
    Cleans the input DataFrame, separates NET and NETPACKET data,
    and creates stack plots for network usage data.

    Parameters:
    - net_df (DataFrame): Input DataFrame containing network and NETPACKET data.
    """
    global CONFIG
    # Debug: Print original shape

    # Remove the header row for NETPACKET
    net_df = net_df[~net_df["Timestamp"].str.contains("Network Packets", na=False)]

    # Split into NET and NETPACKET
    net_data = net_df[net_df["NET"] == "NET"].copy()
    netpacket_data = net_df[net_df["NET"] == "NETPACKET"].copy()

    start_datetime = datetime.strptime(start_time, "%H:%M:%S")
    net_data['Timestamp'] = [
        start_datetime + timedelta(seconds=i * interval_seconds) for i in range(num_records)
    ]

    start_datetime = datetime.strptime(start_time, "%H:%M:%S")
    netpacket_data['Timestamp'] = [
        start_datetime + timedelta(seconds=i * interval_seconds) for i in range(num_records)
    ]

    # Dynamically detect columns to convert
    columns_to_convert = [col for col in net_df.columns if 'KB/s' in col]

    # Convert numeric columns to float
    for col in columns_to_convert:
        net_data[col] = pd.to_numeric(net_data[col], errors="coerce")
        netpacket_data[col] = pd.to_numeric(netpacket_data[col], errors="coerce")

    # Plot NET data dynamically
    plt.figure(figsize=(18, 6))
    net_data_columns = [col for col in net_data.columns if 'KB/s' in col]
    plt.stackplot(
        net_data["Timestamp"],
        [net_data[col] for col in net_data_columns],
        labels=net_data_columns
    )
    plt.title(f'Network Traffic (NET)  {CONFIG["HOSTNAME"]} - {CONFIG["DATE"]}')
    plt.xlabel("Timestamp")
    plt.ylabel("KB/s")
    plt.legend(loc="upper left")

    # Set x-axis ticks to 30-minute intervals
    x_ticks = pd.date_range(start=start_datetime, periods=num_records // 60, freq="30min")
    plt.xticks(x_ticks, [t.strftime("%H:%M:%S") for t in x_ticks], rotation=45)

    # Save plot as PNG
    plt.tight_layout()
    png_file = os.path.join(CONFIG["HOSTOUTDIR"], f'{CONFIG["HOSTNAME"]}_NETWORK.png')
    plt.savefig(png_file)
    plt.close()

    # Plot NETPACKET data dynamically
    plt.figure(figsize=(18, 6))
    netpacket_data_columns = [col for col in netpacket_data.columns if 'KB/s' in col]
    plt.stackplot(
        netpacket_data["Timestamp"],
        [netpacket_data[col] for col in netpacket_data_columns],
        labels=netpacket_data_columns
    )
    plt.title(f'Network Traffic (NETPACKET)  {CONFIG["HOSTNAME"]} - {CONFIG["DATE"]}')
    plt.xlabel("Timestamp")
    plt.ylabel("KB/s")
    plt.legend(loc="upper left")

    # Set x-axis ticks to 30-minute intervals
    x_ticks = pd.date_range(start=start_datetime, periods=num_records // 60, freq="30min")
    plt.xticks(x_ticks, [t.strftime("%H:%M:%S") for t in x_ticks], rotation=45)

    # Save plot as PNG
    plt.tight_layout()
    png_file = os.path.join(CONFIG["HOSTOUTDIR"], f'{CONFIG["HOSTNAME"]}_NETPACKETS.png')
    plt.savefig(png_file)
    plt.close()


def process_and_plot_top_data(top_df, top_n=80):
    """
    Processes the TOP data, calculates resource usage for each command, and plots the results
    for the top N commands with dynamic x-axis management.

    Parameters:
    - top_df (DataFrame): The input TOP DataFrame.
    - top_n (int): The number of top commands to display based on average CPU usage.
    """
    global CONFIG
    # Ensure numeric columns are converted properly
    numeric_cols = ["%CPU", "%Usr", "%Sys", "Size", "ResSet", "IOwaitTime"]
    for col in numeric_cols:
        top_df[col] = pd.to_numeric(top_df[col], errors="coerce")

    # Drop rows with missing or invalid values in critical columns
    top_df.dropna(subset=["%CPU", "Size", "IOwaitTime"], inplace=True)

    # Group by Command and calculate statistics
    grouped = top_df.groupby("Command").agg(
        Avg_CPU=("%CPU", "mean"),
        Max_CPU=("%CPU", "max"),
        Min_Memory=("Size", "min"),
        Avg_Memory=("Size", "mean"),
        Max_Memory=("Size", "max"),
        Avg_IO=("IOwaitTime", "mean"),
        Max_IO=("IOwaitTime", "max"),
    ).reset_index()

    # Sort by Average CPU and limit to top N commands
    grouped = grouped.sort_values(by="Avg_CPU", ascending=False).head(top_n)

    # Generate x-axis indices
    x_ticks = grouped["Command"]
    x_indices = np.arange(len(x_ticks))  # Indexes for the x-axis

    # Plot CPU usage
    plt.figure(figsize=(14, 6))
    plt.bar(x_indices, grouped["Avg_CPU"], label="Average CPU", color="blue", alpha=0.7)
    plt.bar(x_indices, grouped["Max_CPU"], label="Max CPU", color="orange", alpha=0.7)
    plt.title(f'CPU Usage by Command  {CONFIG["HOSTNAME"]} - {CONFIG["DATE"]}')
    plt.xlabel("Command")
    plt.ylabel("%CPU")
    plt.xticks(x_indices, x_ticks, rotation=45, ha="right")  # Dynamically set ticks
    plt.legend()

    # Save plot as PNG
    plt.tight_layout()
    png_file = os.path.join(CONFIG["HOSTOUTDIR"], f'{CONFIG["HOSTNAME"]}_CPUBYCMD.png')
    plt.savefig(png_file)
    plt.close()

    # Plot Memory usage
    plt.figure(figsize=(14, 6))
    plt.bar(x_indices, grouped["Avg_Memory"], label="Average Memory", color="blue", alpha=0.7)
    plt.bar(x_indices, grouped["Max_Memory"], label="Max Memory", color="orange", alpha=0.7)
    plt.title(f'Memory Usage by Command  {CONFIG["HOSTNAME"]} - {CONFIG["DATE"]}')
    plt.xlabel("Command")
    plt.ylabel("Memory Size")
    plt.xticks(x_indices, x_ticks, rotation=45, ha="right")  # Dynamically set ticks
    plt.legend()

    # Save plot as PNG
    plt.tight_layout()
    png_file = os.path.join(CONFIG["HOSTOUTDIR"], f'{CONFIG["HOSTNAME"]}_MEMBYCMD.png')
    plt.savefig(png_file)
    plt.close()

    # Plot I/O usage
    plt.figure(figsize=(14, 6))
    plt.bar(x_indices, grouped["Avg_IO"], label="Average I/O", color="blue", alpha=0.7)
    plt.bar(x_indices, grouped["Max_IO"], label="Max I/O", color="orange", alpha=0.7)
    plt.title(f'I/O Usage by Command {CONFIG["HOSTNAME"]} - {CONFIG["DATE"]}')
    plt.xlabel("Command")
    plt.ylabel("I/O Wait Time")
    plt.xticks(x_indices, x_ticks, rotation=45, ha="right")  # Dynamically set ticks
    plt.legend()

    # Save plot as PNG
    plt.tight_layout()
    png_file = os.path.join(CONFIG["HOSTOUTDIR"], f'{CONFIG["HOSTNAME"]}_IOBYCMD.png')
    plt.savefig(png_file)
    plt.close()


def create_document_with_images(metrics, image_dir, output_dir, document_name, hostname, date, image_height=3, image_width=7, main_heading="NMON Analysis Report"):
    """
    Create a Word document with images added two per page, with proper formatting.
    """
    doc = Document()
    
    # Add styles for consistent formatting
    styles = doc.styles
    style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
    style.font.size = Pt(16)
    style.font.bold = True
    style.paragraph_format.space_after = Pt(12)
    
    # Add main heading with custom style
    doc.add_heading(main_heading, level=1).style = 'CustomHeading1'
    
    for metric in metrics:
        matched_images = [
            os.path.join(image_dir, img)
            for img in os.listdir(image_dir)
            if metric.lower() in img.lower() and img.lower().endswith('.png')
        ]
        
        if not matched_images:
            continue
            
        # Add section heading
        section_heading = doc.add_heading(f'{metric} Usage on server {hostname} for the date {date}', level=2)
        section_heading.style = 'CustomHeading1'
        
        # Add images with proper sizing and positioning
        for idx, image_path in enumerate(matched_images):
            try:
                # Add paragraph for Continuing the code exactly where it left off:

                # Add paragraph for image container
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                
                # Add image with specific dimensions
                run.add_picture(image_path, width=Inches(image_width), height=Inches(image_height))
                
                # Add spacing after image
                p.paragraph_format.space_after = Pt(12)
                
                # Add page break after every two images
                if (idx + 1) % 2 == 0:
                    doc.add_page_break()
                    
            except Exception as e:
                logging.error(f"Error adding image {image_path}: {e}")
                
    # Save document
    doc_path = os.path.join(output_dir, document_name)
    doc.save(doc_path)
    return doc_path


def concat_documents_with_styles(top_docx, middle_docx, bottom_docx, output_docx):
    """
    Merge documents while preserving formatting and properly handling headers/footers.
    """
    # Create composite document
    composite = Document()
    
    # Copy styles from source documents
    for doc_path in [top_docx, middle_docx, bottom_docx]:
        source = Document(doc_path)
        for style in source.styles:
            if style.name not in composite.styles:
                composite.styles.add_style(style.name, style.type)
    
    # Process top document
    top_doc = Document(top_docx)
    header, footer = extract_header_footer(top_doc)
    
    # Set up sections
    composite.sections[0].different_first_page_header_footer = True
    composite.sections[0].header.is_linked_to_previous = False
    composite.sections[0].footer.is_linked_to_previous = False
    
    # Copy header/footer content
    copy_section_format(header, composite.sections[0].header)
    copy_section_format(footer, composite.sections[0].footer)
    
    # Copy content from each document
    for source_doc in [top_doc, Document(middle_docx), Document(bottom_docx)]:
        # Add section break
        if composite.paragraphs:
            composite.add_section()
        
        # Copy content
        for element in source_doc.element.body:
            composite.element.body.append(element)
            
        # Apply header/footer to new section
        if composite.sections:
            section = composite.sections[-1]
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            copy_section_format(header, section.header)
            copy_section_format(footer, section.footer)
    
    # Save final document
    composite.save(output_docx)


def copy_section_format(source, target):
    """
    Copy formatting from source section to target section.
    """
    for paragraph in source.paragraphs:
        new_paragraph = target.add_paragraph()
        for run in paragraph.runs:
            new_run = new_paragraph.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            new_run.font.name = run.font.name
            if run.font.size:
                new_run.font.size = run.font.size
            if run._element.rPr.color is not None:
                new_run.font.color.rgb = run.font.color.rgb


def extract_header_footer(doc):
    """
    Extract header and footer with proper formatting.
    """
    header = doc.sections[0].header
    footer = doc.sections[0].footer
    
    # Ensure we're getting all content including tables and images
    header_content = Document()
    header_content.add_section()
    for element in header.element:
        header_content.element.body.append(element)
        
    footer_content = Document()
    footer_content.add_section()
    for element in footer.element:
        footer_content.element.body.append(element)
        
    return header_content, footer_content


def send_email(host_dirs, body, out_dir):
    """
    Send an email with all the generated documents as attachments.

    Args:
        smtp_server (str): SMTP server address (e.g., 'localhost' for Sendmail).
        receiver_email (str): Email address of the receiver.
        subject (str): Email subject line.
        body (str): Email body text.
        host_dirs (list): List of host-specific directories.
        out_dir (str): Parent directory containing host-specific directories.
    """
    global CONFIG
    receiver_email = "receiver_mail_here"
    # Create the email message
    msg = MIMEMultipart()
    msg['From'] = 'sender_mail_here'  # Change as needed
    msg['To'] = receiver_email
    msg['Subject'] = "NMON Analysis Reports"

    # Add the body text
    msg.attach(MIMEText(body, 'plain'))

    # Iterate through each host's directory to attach documents
    for host_dir in host_dirs:
        if os.path.isdir(out_dir):
            for file_name in os.listdir(out_dir):
                if file_name.endswith("diagnostic_v1.0.docx"):
                    file_path = os.path.join(out_dir, file_name)
                    with open(file_path, "rb") as attachment:
                        # Add the file as an email attachment
                        part = MIMEBase("application", "octet-stream")
                        part.set_payload(attachment.read())
                        encoders.encode_base64(part)
                        part.add_header(
                            "Content-Disposition",
                            f"attachment; filename={file_name}",
                        )
                        msg.attach(part)

    # Send the email
    try:
        with smtplib.SMTP(CONFIG["SMTPSERVER"]) as server:
            server.sendmail(msg['From'], receiver_email, msg.as_string())
        print(f"Email sent successfully to {receiver_email}")
    except Exception as e:
        print(f"Error sending email: {e}")


def main():
    host_dirs = []  # Keep track of host directories created

    # Iterate through all combinations of hosts and dates
    for host, date in product(hosts, dates):
        CONFIG["HOSTNAME"] = host
        CONFIG["DATE"] = date
        CONFIG["HOSTOUTDIR"] = f'{OUTDIR}/{host}'
        CONFIG["DOCUMENT_NAME"] = f'{host}_{date}_nmon.docx'
        CONFIG["FINALDOCNAME"] = f'{CONFIG["HOSTOUTDIR"]}/{host}_{date}_Formulaire_de_diagnostic_v1.0.docx'

        os.makedirs(CONFIG["HOSTOUTDIR"], exist_ok=True)
        remote_path = f"{CONFIG['REMOTE_DIR']}/{date}.nmon"
        file_path = f"{CONFIG['OUTDIR']}/{date}.nmon"

        # Replace 'hostname' and 'user' with actual values as needed
        ssh_copy_file(CONFIG["HOSTNAME"], remote_path, file_path)

        dataframes = process_system_data(file_path)
        top_df = dataframes.get("TOP")
        cpu_df = dataframes.get("CPU_ALL")
        mem_df = dataframes.get("MEM")
        vm_df = dataframes.get("VM")
        net_df = dataframes.get("NET")
        fs_df = dataframes.get("JFSFILE")
        diskb_df = dataframes.get("DISKBUSY")

        # Call plotting functions
        plot_cpu(cpu_df)
        plot_memory(mem_df)
        plot_swap(mem_df)
        plot_filesystem(fs_df)
        plot_network_usage(net_df)
        process_and_plot_top_data(top_df)

        # Generate document with images
        middle_doc_path = create_document_with_images(
            CONFIG["METRICS"],
            CONFIG["HOSTOUTDIR"],
            CONFIG["HOSTOUTDIR"],
            CONFIG["DOCUMENT_NAME"],
            CONFIG["HOSTNAME"],
            CONFIG["DATE"],
        )

        concat_documents_with_styles(CONFIG["TOPDOCFORMAT"], middle_doc_path, CONFIG["BOTTOMDOCFORMAT"], CONFIG["FINALDOCNAME"])

        # Track created host directory
        if CONFIG["HOSTOUTDIR"] not in host_dirs:
            host_dirs.append(CONFIG["HOSTOUTDIR"])

    # Send email with all documents
    email_body = (
        "Attached are the NMON analysis reports for the requested hosts and dates. "
        "Please review the documents."
    )
    send_email(host_dirs, email_body, CONFIG["HOSTOUTDIR"])


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Analyze NMON file and send report.")
    parser.add_argument('--host', required=True, help="Comma-separated list of target hosts (e.g., host1,host2).")
    parser.add_argument('--date', required=True, help="Comma-separated list of dates (e.g., 2024-11-21,2024-11-22).")
    args = parser.parse_args()

    # Parse hosts and dates into lists
    hosts = args.host.split(',')
    dates = args.date.split(',')

    # Output directory
    PREREQDOC= "./PREREQDOC"
    OUTDIR = "./output"
    os.makedirs(OUTDIR, exist_ok=True)

    CONFIG = {
            "OUTDIR": OUTDIR,
            "REMOTE_DIR": "/tools/list/admin/perf/nmon",  # Directory on the remote server
            "METRICS": ['CPU', 'MEM', 'NET', 'FileSystem', 'SWAP', 'IO'],
            "SMTPSERVER": "smtp.example.com",
            "TOPDOCFORMAT": f'{PREREQDOC}/top_doc_format.docx',
            "BOTTOMDOCFORMAT": f'{PREREQDOC}/bottom_doc_format.docx',
        }
    main()
