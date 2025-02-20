from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def concatenate_documents(top_file, middle_file, bottom_file, output_file):
    """
    Concatenates three Word documents into one, maintaining formatting.
    Uses header/footer from top document across all sections.
    
    Parameters:
    - top_file (str): Path to the top document
    - middle_file (str): Path to the middle document
    - bottom_file (str): Path to the bottom document
    - output_file (str): Path for the final concatenated document
    """
    # Load all three documents
    top_doc = Document(top_file)
    middle_doc = Document(middle_file)
    bottom_doc = Document(bottom_file)

    # Create a new document for the output
    final_doc = Document()

    # Helper function to copy header/footer from source to target section
    def copy_header_footer(source_section, target_section):
        # Copy header
        for header in source_section.header_tables:
            new_header = target_section.header
            for paragraph in header.paragraphs:
                new_paragraph = new_header.add_paragraph()
                new_paragraph.text = paragraph.text
                new_paragraph.style = paragraph.style
        
        # Copy footer
        for footer in source_section.footer_tables:
            new_footer = target_section.footer
            for paragraph in footer.paragraphs:
                new_paragraph = new_footer.add_paragraph()
                new_paragraph.text = paragraph.text
                new_paragraph.style = paragraph.style

    # Copy content from top document
    for element in top_doc.element.body:
        final_doc.element.body.append(element)

    # Add a section break before middle content
    final_doc.add_section()

    # Copy content from middle document
    for element in middle_doc.element.body:
        final_doc.element.body.append(element)

    # Add a section break before bottom content
    final_doc.add_section()

    # Copy content from bottom document
    for element in bottom_doc.element.body:
        final_doc.element.body.append(element)

    # Apply top document's header/footer to all sections in final document
    top_first_section = top_doc.sections[0]
    for section in final_doc.sections:
        copy_header_footer(top_first_section, section)
        
        # Preserve section properties (margins, orientation, etc.)
        section.top_margin = top_first_section.top_margin
        section.bottom_margin = top_first_section.bottom_margin
        section.left_margin = top_first_section.left_margin
        section.right_margin = top_first_section.right_margin
        section.page_height = top_first_section.page_height
        section.page_width = top_first_section.page_width

    # Save the final document
    final_doc.save(output_file)
    print(f"Documents concatenated successfully. Saved as: {output_file}")

# Example usage
if __name__ == "__main__":
    # Replace these with your actual file paths
    top_file = "top_document.docx"
    middle_file = "middle_document.docx"
    bottom_file = "bottom_document.docx"
    output_file = "final_document.docx"

    try:
        concatenate_documents(top_file, middle_file, bottom_file, output_file)
    except FileNotFoundError as e:
        print(f"Error: One or more input files not found - {e}")
    except Exception as e:
        print(f"An error occurred: {e}")
