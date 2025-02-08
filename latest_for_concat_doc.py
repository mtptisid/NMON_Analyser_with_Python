from docx import Document
from docx.oxml import parse_xml
from copy import deepcopy
from io import BytesIO

def concat_documents_with_styles(top_file, middle_file, bottom_file, output_file):
    """
    Merges documents while preserving top document's headers/footers,
    styles from all documents, and properly handling images
    """
    # Load documents
    top_doc = Document(top_file)
    middle_doc = Document(middle_file)
    bottom_doc = Document(bottom_file)

    # Merge styles from all documents
    merge_styles(middle_doc, top_doc)
    merge_styles(bottom_doc, top_doc)

    # Append content with proper image handling
    append_doc_content(top_doc, middle_doc)
    append_doc_content(top_doc, bottom_doc)

    # Save merged document
    top_doc.save(output_file)

def append_doc_content(target_doc, source_doc):
    """Appends content from source to target document with image handling"""
    # Copy all elements from source to target
    for element in source_doc.element.body:
        # Deep copy preserves original element structure
        element_copy = deepcopy(element)
        target_doc.element.body.append(element_copy)
        
        # Process images in the copied element
        process_images_in_element(element_copy, source_doc, target_doc)

def process_images_in_element(element, source_doc, target_doc):
    """Handles image relationships in copied elements"""
    # Define namespace mappings
    namespaces = {
        'wp': "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        'a': "http://schemas.openxmlformats.org/drawingml/2006/main",
        'r': "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    }

    # Find all image elements
    for inline in element.xpath('.//wp:inline', namespaces=namespaces):
        blip = inline.xpath('.//a:blip', namespaces=namespaces)
        if not blip:
            continue
        blip = blip[0]
        embed_id = blip.get('{{{}}}embed'.format(namespaces['r']))
        
        if embed_id and embed_id in source_doc.part.related_parts:
            # Get image data from source document
            image_part = source_doc.part.related_parts[embed_id]
            image_bytes = BytesIO(image_part.blob)
            
            # Add image to target document and get new relationship ID
            new_embed_id = target_doc.part.relate_to(image_part, image_part.content_type)
            
            # Update the embed ID in the copied element
            blip.set('{{{}}}embed'.format(namespaces['r']), new_embed_id)

def merge_styles(source_doc, target_doc):
    """Merges styles from source document to target document"""
    for style in source_doc.styles:
        if not target_doc.styles.get(style.name, None):
            # Add missing style to target document
            new_style = target_doc.styles.add_style(style.name, style.type)
            new_style.element = deepcopy(style.element)
